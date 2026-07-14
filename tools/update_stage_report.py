from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Rapport_de_stage_StageLink_modifie.docx"
ASSET_DIR = ROOT / "rapport_assets"
ASSET_DIR.mkdir(exist_ok=True)


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


TITLE = font(34, True)
HEAD = font(22, True)
TEXT = font(17)
SMALL = font(15)


def draw_wrapped(draw, xy, text, fnt, fill=(31, 41, 55), max_width=280, line_gap=5):
    words = text.split()
    lines, current = [], ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def box(draw, xy, title, items, width=310, height=150, fill="#FFFFFF", outline="#2563EB"):
    x, y = xy
    draw.rounded_rectangle((x, y, x + width, y + height), radius=16, fill=fill, outline=outline, width=3)
    draw.rounded_rectangle((x, y, x + width, y + 42), radius=16, fill=outline, outline=outline)
    draw.rectangle((x, y + 26, x + width, y + 42), fill=outline)
    draw.text((x + 16, y + 9), title, font=HEAD, fill="white")
    yy = y + 54
    for item in items:
        draw.text((x + 16, yy), item, font=SMALL, fill="#111827")
        yy += 23


def connector(draw, p1, p2, label="", color="#64748B"):
    draw.line((p1[0], p1[1], p2[0], p2[1]), fill=color, width=3)
    r = 5
    draw.ellipse((p1[0]-r, p1[1]-r, p1[0]+r, p1[1]+r), fill=color)
    draw.ellipse((p2[0]-r, p2[1]-r, p2[0]+r, p2[1]+r), fill=color)
    if label:
        mx, my = (p1[0] + p2[0]) // 2, (p1[1] + p2[1]) // 2
        bbox = draw.textbbox((0, 0), label, font=SMALL)
        draw.rounded_rectangle((mx - 8, my - 18, mx + (bbox[2]-bbox[0]) + 8, my + 8), radius=7, fill="#F8FAFC", outline="#CBD5E1")
        draw.text((mx, my - 14), label, font=SMALL, fill="#334155")


def save_mcd(path):
    img = Image.new("RGB", (1800, 1240), "#F8FAFC")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "MCD - Modele Conceptuel des Donnees StageLink", font=TITLE, fill="#0F172A")
    d.text((60, 88), "Vue conceptuelle des acteurs, candidatures, offres, entretiens, messages et notifications.", font=TEXT, fill="#475569")

    box(d, (80, 160), "UTILISATEUR", ["id", "nom, prenom", "display_name", "email", "role", "photo", "settings", "2FA"], 330, 215, outline="#1D4ED8")
    box(d, (520, 160), "CANDIDAT", ["id", "user_id", "profil", "competences", "cv/photo", "documents"], 310, 185, outline="#059669")
    box(d, (980, 160), "RH", ["id", "user_id", "entreprise_id", "fonction", "is_admin"], 300, 170, outline="#7C3AED")
    box(d, (1410, 160), "ENTREPRISE", ["id", "nom", "logo", "secteur", "taille", "coordonnees"], 300, 170, outline="#EA580C")

    box(d, (220, 520), "CANDIDATURE", ["id", "candidat", "entreprise", "offre", "statut", "score", "commentaire"], 330, 190, outline="#0EA5E9")
    box(d, (760, 520), "OFFRE", ["id", "entreprise", "rh createur", "titre", "duree", "statut"], 300, 170, outline="#DC2626")
    box(d, (1230, 520), "ENTRETIEN", ["id", "candidature", "recruteur", "date", "heure", "lieu/lien", "statut"], 320, 195, outline="#9333EA")

    box(d, (430, 880), "MESSAGE", ["id", "expediteur", "destinataire", "contenu", "piece jointe", "read_at"], 330, 170, outline="#0891B2")
    box(d, (1020, 880), "NOTIFICATION", ["id", "candidat", "message", "date_envoi", "lu"], 330, 155, outline="#16A34A")

    connector(d, (410, 250), (520, 250), "1,1 / 0,1")
    connector(d, (1280, 245), (1410, 245), "0,n / 1,1")
    connector(d, (830, 250), (980, 250), "1,1 / 0,1")
    connector(d, (520, 620), (410, 300), "depose")
    connector(d, (550, 610), (760, 610), "concerne")
    connector(d, (910, 520), (1440, 330), "publiee par")
    connector(d, (550, 690), (1230, 620), "planifie")
    connector(d, (760, 965), (1020, 965), "genere")
    connector(d, (410, 345), (430, 930), "envoie/recoit")

    img.save(path, quality=95)


def save_mld(path):
    img = Image.new("RGB", (1800, 1300), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "MLD - Modele Logique des Donnees StageLink", font=TITLE, fill="#0F172A")
    d.text((60, 88), "Relations principales avec cles primaires (PK) et cles etrangeres (FK).", font=TEXT, fill="#475569")
    tables = [
        ((70, 150), "users", ["PK id", "nom, prenom, display_name", "email UNIQUE", "role", "password", "photo", "settings JSON", "two_factor_*"]),
        ((520, 150), "entreprises", ["PK id", "nom, adresse", "telephone, email", "logo", "site_web", "secteur_activite", "taille"]),
        ((970, 150), "r_h_s", ["PK id", "FK user_id -> users", "FK entreprise_id -> entreprises", "fonction", "is_admin"]),
        ((1370, 150), "candidats", ["PK id", "FK user_id -> users", "telephone", "filiere, niveau", "photo, cv", "competences"]),
        ((90, 530), "offres", ["PK id", "FK entreprise_id", "FK r_h_id", "titre", "type_stage", "duree", "statut"]),
        ((520, 530), "candidatures", ["PK id", "FK candidat_id", "FK entreprise_id", "FK r_h_id", "FK offre_id NULL", "statut, score", "documents"]),
        ((970, 530), "entretiens", ["PK id", "FK candidature_id", "FK recruteur_id -> users", "date_entretien", "heure", "lieu, type, statut"]),
        ((1370, 530), "messages", ["PK id", "FK sender_id -> users", "FK receiver_id -> users", "body", "attachment", "read_at"]),
        ((520, 900), "candidate_documents", ["PK id", "FK candidat_id", "nom_fichier", "type_document", "chemin", "is_default"]),
        ((970, 900), "notifications", ["PK id", "FK candidat_id", "message", "date_envoi", "lu"]),
    ]
    centers = {}
    for xy, title, items in tables:
        box(d, xy, title, items, width=350, height=230 if len(items) > 6 else 205, outline="#2563EB")
        centers[title] = (xy[0] + 175, xy[1] + 105)
    connector(d, (420, 255), (970, 255), "users 1,n RH")
    connector(d, (870, 255), (970, 255), "")
    connector(d, (1370, 255), (420, 255), "users 1,1 candidat")
    connector(d, (870, 255), (970, 255), "")
    connector(d, (870, 250), (970, 250), "entreprise/RH")
    connector(d, (870, 640), (970, 640), "candidature/entretien")
    connector(d, (870, 1010), (970, 1010), "candidat/notif")
    connector(d, (700, 900), (700, 735), "documents")
    connector(d, (440, 640), (520, 640), "offre")
    connector(d, (1545, 380), (1545, 530), "messages")
    img.save(path, quality=95)


def save_mpd(path):
    img = Image.new("RGB", (1800, 1240), "#F8FAFC")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "MPD - Modele Physique des Donnees StageLink", font=TITLE, fill="#0F172A")
    d.text((60, 88), "Extrait SQL simplifie, aligne avec les migrations Laravel actuelles.", font=TEXT, fill="#475569")
    sql_blocks = [
        ("users", "id BIGINT PK\nnom VARCHAR(255)\nprenom VARCHAR(255)\ndisplay_name VARCHAR(255) NULL\nemail VARCHAR(255) UNIQUE\nrole VARCHAR(255)\npassword VARCHAR(255)\nphoto VARCHAR(255) NULL\nsettings JSON NULL\ntwo_factor_enabled_at TIMESTAMP NULL\ntwo_factor_method VARCHAR(255) NULL\ntwo_factor_code VARCHAR(255) NULL\ntwo_factor_expires_at TIMESTAMP NULL"),
        ("candidatures", "id BIGINT PK\ncandidat_id BIGINT FK\nentreprise_id BIGINT FK\nr_h_id BIGINT FK\noffre_id BIGINT FK NULL\ndate_candidature DATE\ntype_stage VARCHAR(255)\nduree VARCHAR(255)\nstatut VARCHAR(255)\nscore DECIMAL(5,2) NULL\ncommentaire_rh TEXT NULL"),
        ("entretiens", "id BIGINT PK\ncandidature_id BIGINT FK\nrecruteur_id BIGINT FK users.id\ndate_entretien DATE\nheure TIME\nlieu VARCHAR(255)\ntype VARCHAR(255)\nstatut VARCHAR(255)\ncommentaires TEXT NULL"),
        ("messages", "id BIGINT PK\nsender_id BIGINT FK users.id\nreceiver_id BIGINT FK users.id\nbody TEXT\nattachment VARCHAR(255) NULL\nread_at TIMESTAMP NULL\ncreated_at / updated_at"),
        ("candidate_documents", "id BIGINT PK\ncandidat_id BIGINT FK\nnom_fichier VARCHAR(255)\ntype_document VARCHAR(255)\nchemin VARCHAR(255)\nis_default BOOLEAN"),
        ("notifications", "id BIGINT PK\ncandidat_id BIGINT FK\nmessage TEXT\ndate_envoi DATETIME\nlu BOOLEAN\ncreated_at / updated_at"),
    ]
    positions = [(70, 150), (650, 150), (1230, 150), (70, 700), (650, 700), (1230, 700)]
    for (title, content), pos in zip(sql_blocks, positions):
        box(d, pos, title, [], width=500, height=420, outline="#0F766E")
        draw_wrapped(d, (pos[0] + 20, pos[1] + 60), content.replace("\n", "   |   "), SMALL, max_width=460, line_gap=7)
    img.save(path, quality=95)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(anchor, text="", style=None):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    new = Paragraph(new_p, anchor._parent)
    if text:
        new.add_run(text)
    if style:
        new.style = style
    return new


def add_bullet(anchor, text):
    p = insert_after(anchor, f"- {text}", "List Paragraph")
    return p


def rebuild_merise_section():
    doc = Document(DOCX_PATH)

    # Small textual updates outside the enterprise section and before the final screenshots.
    replacements = {
        "dashboards RH_ADMIN et RH_USER": "dashboards RH_ADMIN, Assistant(e) et Stagiaire",
        "acteurs (candidat, RH_ADMIN, RH_USER)": "acteurs (candidat, RH_ADMIN, Assistant(e), Stagiaire)",
        "dashboard RH_ADMIN, dashboard RH_USER": "dashboard RH_ADMIN, dashboard Assistant(e) et dashboard Stagiaire",
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                for run in p.runs:
                    run.text = run.text.replace(old, new)

    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("III.3.1."))
    end = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("III.3.4."))
    anchor = doc.paragraphs[start - 1]
    for p in list(doc.paragraphs[start:end]):
        remove_paragraph(p)

    mcd = ASSET_DIR / "mcd_stagelink.png"
    mld = ASSET_DIR / "mld_stagelink.png"
    mpd = ASSET_DIR / "mpd_stagelink.png"
    save_mcd(mcd)
    save_mld(mld)
    save_mpd(mpd)

    p = insert_after(anchor, "III.3.1. Etude d'opportunite", "Heading 3")
    p = insert_after(p, "L'etude d'opportunite a permis de confirmer le besoin d'une plateforme centralisee pour fluidifier la recherche de stage et le suivi des candidatures. StageLink repond a un probleme concret : les candidats doivent pouvoir postuler, suivre leurs dossiers, recevoir des notifications et echanger avec les recruteurs, tandis que les equipes RH doivent publier des offres, analyser les candidatures, planifier les entretiens et collaborer selon des permissions adaptees.")
    p = insert_after(p, "Acteurs concernes", "Heading 4")
    for item in [
        "Candidat : cree son compte, complete son profil, depose ses documents, postule aux offres, suit ses candidatures, recoit des notifications et communique avec le RH.",
        "RH_ADMIN : gere l'entreprise, les offres, les candidatures, les entretiens, les messages, les notifications et les membres de l'equipe RH.",
        "Assistant(e) : consulte les candidatures, les CV et les lettres de motivation, puis peut valider une candidature selon les droits accordes.",
        "Stagiaire : consulte les candidatures et les documents, sans pouvoir valider ni planifier d'entretien.",
        "Entreprise : porte les offres de stage, les informations institutionnelles et le logo affiche dans l'application.",
    ]:
        p = add_bullet(p, item)
    p = insert_after(p, "Besoins fonctionnels actualises", "Heading 4")
    for item in [
        "Authentification securisee avec double authentification par e-mail, extensible plus tard vers SMS et QR code.",
        "Gestion dynamique des profils, photos, parametres, langues, notifications et preferences utilisateur.",
        "Publication et suivi des offres de stage rattachees a une entreprise.",
        "Depot et suivi des candidatures avec statut, score, documents et commentaires RH.",
        "Messagerie directe entre candidats et RH avec historique, pieces jointes et compteur de lecture.",
        "Planification des entretiens avec date, heure, recruteur, lieu ou lien, statut et commentaires.",
        "Gestion d'equipe avec roles RH_ADMIN, Assistant(e) et Stagiaire, chacun disposant d'un tableau de bord adapte.",
    ]:
        p = add_bullet(p, item)

    p = insert_after(p, "III.3.2. Analyse fonctionnelle", "Heading 3")
    p = insert_after(p, "La plateforme est organisee autour de modules fonctionnels relies entre eux : authentification, espace candidat, espace RH, gestion des offres, candidatures, entretiens, messagerie, notifications, parametres et gestion de l'entreprise. Cette organisation permet de separer les responsabilites tout en conservant une experience utilisateur coherente.")
    p = insert_after(p, "Fonctionnalites principales", "Heading 4")
    for item in [
        "Accueil, inscription et connexion avec orientation automatique vers le tableau de bord correspondant au role.",
        "Espace candidat : profil, documents, offres disponibles, candidatures, messages, notifications, securite et langue.",
        "Espace RH_ADMIN : entreprise, offres, candidatures, entretiens, equipe, messages, notifications et parametres.",
        "Espace Assistant(e) : consultation des candidatures, documents et entretiens, avec validation autorisee.",
        "Espace Stagiaire : consultation des candidatures, documents et entretiens en lecture seule.",
        "Systeme de notification automatique lors des candidatures, changements de statut, entretiens et messages.",
    ]:
        p = add_bullet(p, item)
    p = insert_after(p, "Cas d'utilisation synthetiques", "Heading 4")
    for item in [
        "Un candidat s'inscrit, complete son profil avec une photo et un CV, puis postule a une offre.",
        "Le RH_ADMIN consulte les candidatures, valide ou refuse un dossier et peut programmer un entretien.",
        "Un candidat envoie un message au RH ; le message apparait dans la messagerie RH et reste conserve dans l'historique.",
        "Un RH_ADMIN ajoute un collaborateur Assistant(e) ou Stagiaire ; a la connexion, celui-ci accede au tableau de bord correspondant a ses droits.",
        "Un utilisateur active la 2FA par e-mail ; lors de la connexion suivante, un code temporaire est envoye avant l'acces au dashboard.",
        "Un utilisateur choisit Francais ou English ; la preference est sauvegardee et restauree automatiquement.",
    ]:
        p = add_bullet(p, item)

    p = insert_after(p, "III.3.3. Conception du systeme (Methode MERISE)", "Heading 3")
    p = insert_after(p, "La methode MERISE a ete utilisee pour structurer les donnees de StageLink en trois niveaux : le modele conceptuel des donnees (MCD), le modele logique des donnees (MLD) et le modele physique des donnees (MPD). Les modeles ci-dessous tiennent compte des dernieres evolutions de l'application : messagerie, entretiens, notifications dynamiques, documents candidats, parametres, langue et double authentification.")
    p = insert_after(p, "a. Regles de gestion", "Heading 4")
    for item in [
        "Un utilisateur possede un seul role principal : candidat, rh_admin, assistant ou stagiaire.",
        "Un candidat est rattache a un compte utilisateur et peut deposer plusieurs documents.",
        "Une entreprise peut publier plusieurs offres et posseder plusieurs membres RH.",
        "Une candidature appartient a un candidat, a une entreprise et peut etre rattachee a une offre.",
        "Une candidature peut donner lieu a un ou plusieurs entretiens planifies.",
        "Un message relie toujours un expediteur et un destinataire, tous deux issus de la table users.",
        "Les notifications candidat sont liees au compte candidat et indiquent si elles ont ete lues.",
        "La 2FA conserve uniquement les informations necessaires : activation, methode, code temporaire et expiration.",
    ]:
        p = add_bullet(p, item)

    p = insert_after(p, "b. Modele Conceptuel des Donnees (MCD)", "Heading 4")
    p = insert_after(p, "Le MCD represente les principales entites de StageLink et leurs associations : utilisateur, candidat, RH, entreprise, offre, candidature, entretien, message, notification et document.")
    p = insert_after(p, "Figure 1 : MCD actualise de StageLink")
    p.alignment = 1
    p = insert_after(p)
    p.alignment = 1
    run = p.add_run()
    run.add_picture(str(mcd), width=Inches(6.6))

    p = insert_after(p, "c. Modele Logique des Donnees (MLD)", "Heading 4")
    p = insert_after(p, "Le MLD traduit le modele conceptuel en relations. Les cles etrangeres materialisent les liens entre les utilisateurs, candidats, entreprises, RH, offres, candidatures, entretiens, messages et notifications.")
    p = insert_after(p, "Figure 2 : MLD actualise de StageLink")
    p.alignment = 1
    p = insert_after(p)
    p.alignment = 1
    run = p.add_run()
    run.add_picture(str(mld), width=Inches(6.6))

    p = insert_after(p, "d. Modele Physique des Donnees (MPD)", "Heading 4")
    p = insert_after(p, "Le MPD correspond a l'implementation relationnelle dans MySQL via les migrations Laravel. Les tables principales sont creees et modifiees par les migrations afin d'assurer la coherence entre le code applicatif et la structure de la base.")
    p = insert_after(p, "Figure 3 : MPD simplifie de StageLink")
    p.alignment = 1
    p = insert_after(p)
    p.alignment = 1
    run = p.add_run()
    run.add_picture(str(mpd), width=Inches(6.6))

    p = insert_after(p, "Extrait du MPD retenu", "Heading 4")
    for item in [
        "users(id, nom, prenom, display_name, email, role, password, photo, settings, two_factor_enabled_at, two_factor_method, two_factor_code, two_factor_expires_at)",
        "entreprises(id, nom, adresse, telephone, email, logo, site_web, description, secteur_activite, taille)",
        "candidats(id, user_id, date_naissance, telephone, filiere, niveau, photo, cv, competences, experiences, langues, use_default_cv)",
        "r_h_s(id, user_id, entreprise_id, fonction, is_admin)",
        "offres(id, entreprise_id, r_h_id, titre, description, type_stage, duree, filiere_cible, competences_requises, lieu, statut)",
        "candidatures(id, candidat_id, entreprise_id, r_h_id, offre_id, date_candidature, statut, score, commentaire_rh, cv, lettre_motivation, lettre_recommandation)",
        "entretiens(id, candidature_id, recruteur_id, date_entretien, heure, lieu, type, statut, commentaires)",
        "messages(id, sender_id, receiver_id, body, attachment, read_at)",
        "notifications(id, candidat_id, message, date_envoi, lu)",
        "candidate_documents(id, candidat_id, nom_fichier, type_document, chemin, is_default)",
    ]:
        p = add_bullet(p, item)

    p = insert_after(p, "e. Modelisation UML", "Heading 4")
    p = insert_after(p, "En complement de MERISE, l'UML permet de decrire les interactions entre les acteurs et le systeme. Les principaux cas d'utilisation concernent l'inscription, la connexion avec 2FA, la gestion du profil, la publication d'offres, le depot de candidatures, la validation, la messagerie, les notifications et la planification des entretiens.")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("À ce stade du stage") or text.startswith("A ce stade du stage"):
            paragraph.text = "A ce stade du stage, la plateforme StageLink dispose des principaux modules attendus : inscription et connexion, creation automatique des comptes candidat et RH, tableaux de bord par role, gestion des entreprises, offres, candidatures, documents, entretiens, messages, notifications, parametres, langue et double authentification par e-mail."
        elif text.startswith("Le module de score de compatibilité") or text.startswith("Le module de score de compatibilite"):
            paragraph.text = "Les evolutions recentes ont permis de rendre dynamiques les parametres RH et candidat, d'ajouter une messagerie complete entre candidats et recruteurs, de finaliser la planification des entretiens, de renforcer les permissions Assistant(e)/Stagiaire et de mettre en place la preference de langue ainsi que la securite 2FA."
        elif text.startswith("Cette phase de développement progressif") or text.startswith("Cette phase de developpement progressif"):
            paragraph.text = "Ces livrables montrent une application fonctionnelle, coherente et evolutive, dont l'architecture Laravel permet d'ajouter de nouvelles methodes d'authentification, des statistiques avancees ou encore des integrations externes sans remettre en cause la base existante."

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    rebuild_merise_section()
    print(f"Updated {DOCX_PATH}")
